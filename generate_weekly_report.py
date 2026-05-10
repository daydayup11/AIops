from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from copy import deepcopy
import copy

template_path = "docs/周报_代雨童_20260505.docx"
output_path = "docs/周报_代雨童_20260510.docx"

doc = Document(template_path)

# Clear all paragraphs content while preserving structure
# Strategy: copy template, then replace text

new_doc = Document(template_path)

# Helper to set paragraph text preserving style
def set_para_text(para, text):
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = text
    else:
        para.add_run(text)

# Paragraphs in template:
# P0: title
# P1: period
# P2: reporter
# P3: Heading "一、本周工作总结"
# P4: summary paragraph
# P5: Heading "二、核心工作与产出"
# P6-P31: numbered items and bullets
# P32: Heading "三、下周工作计划"
# P33-P36: bullet points

paras = new_doc.paragraphs

# P0: title
set_para_text(paras[0], "项目周报：AIops 校园网络分析平台 开发")

# P1: period
set_para_text(paras[1], "汇报周期：2026年5月5日 - 2026年5月10日")

# P2: reporter (unchanged)
# paras[2] already correct

# P3: heading unchanged

# P4: overall summary
set_para_text(paras[4],
    "本周从零开始独立完成了 AIops 校园网络分析平台（Campus Network Analyzer）的全栈开发。"
    "该平台基于 LangGraph + FastAPI + React 技术栈，允许用户通过自然语言查询 ClickHouse 中的校园网络数据，"
    "由大模型自动生成数据分析 Python 脚本、执行代码并渲染可视化图表。"
    "本周完成了后端 Agent Pipeline 的设计与实现、前端交互界面的开发、可视化管线的多次重构迭代，"
    "以及日志系统、测试体系的完善，项目已具备端到端可运行能力。"
)

# P5: heading unchanged

# Now replace items P6 onwards
# P6: item 1 title
set_para_text(paras[6], "1. 项目从零搭建：完成全栈脚手架与基础设施")

set_para_text(paras[7],
    "完成项目初始化与整体架构设计，包括：Python FastAPI WebSocket 服务端、"
    "React/Vite 前端脚手架（含 Tailwind CSS、shadcn/ui 组件库接入）、"
    "ClickHouse 查询层（含 SQL 安全守卫）、SQLite 会话与消息持久化存储、"
    "以及 Pydantic Schema 体系定义。项目采用 LangGraph 构建 Agent 工作流，"
    "通过 WebSocket 实现前后端实时通信。"
)

set_para_text(paras[8], "2. LangGraph Agent Pipeline 设计与多轮重构")

set_para_text(paras[9],
    "核心 Pipeline 经历了两轮重大架构演进：初版实现 Planner → SQL Engineer → Visualizer 三节点流程；"
    "第二版引入 Clarifier（意图澄清）、Summarizer（分析报告生成）节点，并设计 VizBlueprint 可视化蓝图机制；"
    "最终版将可视化管线重构为脚本化方案（Script-Based Viz Pipeline）："
)

set_para_text(paras[10],
    "·\xa0Planner：解析用户意图，输出 AnalysisPlan（分析计划），不再承担可视化蓝图职责；"
)

set_para_text(paras[11],
    "·\xa0SQL Engineer：生成包含数据查询与 matplotlib/seaborn 可视化的完整 Python 脚本，替代原 SQL Tasks 方案；"
)

set_para_text(paras[12],
    "·\xa0Code Reviewer：对生成脚本进行安全性与性能审查，拦截危险操作；"
)

set_para_text(paras[13],
    "·\xa0Script Runner：在受控子进程中执行 Python 脚本，输出 PNG 图表与 data_summary.json；"
)

set_para_text(paras[14],
    "·\xa0Summarizer：读取图表数量与数据摘要，生成结构化自然语言分析报告，通过 WebSocket 推送前端。"
)

set_para_text(paras[15],
    "重构过程中同步修复了多个关键问题：ClickHouse execute() 返回 DataFrame 的正确处理模式、"
    "LLM 调用期间 WebSocket 心跳保活（防止长查询超时断连）、子进程树的完整清理机制。"
)

set_para_text(paras[16], "3. 前端 UI 全面重设计")

set_para_text(paras[17],
    "引入 shadcn/ui 组件库（Button、Input、ScrollArea、Tooltip），对全部前端组件进行视觉重构："
    "App Header 改为图标按钮+Tooltip 交互，ChatPanel 采用精细化气泡与卡片结果布局，"
    "SessionSidebar 加入时间戳展示，InputBar 与 ProgressBar 统一使用 shadcn 风格。"
    "同时实现暗色科技主题（Tech Dark Theme）：新增 CSS 变量体系、ECharts Tech Theme 注册与自动应用、"
    "useTheme Hook（localStorage 持久化），以及月亮图标主题切换按钮。"
    "前端支持新的 image render type，移除 ECharts 直接渲染与 PlanCard 组件，改为展示后端生成的 PNG 图表。"
)

set_para_text(paras[18], "4. 日志系统建设")

set_para_text(paras[19],
    "为项目全链路引入结构化日志体系：实现 JsonFormatter 与 setup_logging() 统一入口，"
    "在 planner、sql_engineer、clickhouse、parallel executor、pipeline、chat API 等各层添加日志覆盖，"
    "修复文件 Handler 重复注册与日志目录创建问题，应用启动时自动初始化日志系统。"
    "同时为 script_runner 添加脚本内容与输出目录的 debug 日志，便于生产环境诊断。"
)

set_para_text(paras[20], "5. 测试体系完善与重构清理")

set_para_text(paras[21],
    "随架构演进同步维护测试体系：移除因脚本化重构而废弃的 test_sql_safety、VizBlueprint/SubTask 相关测试，"
    "删除 visualizer 与 parallel executor 模块及对应测试，加强 test_schemas 断言覆盖。"
    "新增 script_runner 错误场景测试（验证 stderr 内容捕获）、PipelineState 字段完整性测试。"
    "同时补充 axp 应用名称映射表到 Schema，修正 sql_engineer prompt 中 appid JOIN axp 的约束要求。"
)

# Keep remaining items or replace with blank if template has more
for i in range(22, 32):
    if i < len(paras):
        set_para_text(paras[i], "")

# P32: heading unchanged

# P33-P36: next week plan
set_para_text(paras[33],
    "·\xa0端到端联调验证：本地启动完整服务（ClickHouse + FastAPI + React），"
    "测试自然语言查询→脚本生成→图表渲染的完整链路，覆盖典型查询场景与异常边界；"
)

set_para_text(paras[34],
    "·\xa0AXP 数据接入优化：完善 axp 应用名称映射的 JOIN 逻辑，"
    "确保 sql_engineer 生成的脚本能正确关联 appid 与应用名，提升查询结果可读性；"
)

set_para_text(paras[35],
    "·\xa0Script Runner 健壮性提升：为脚本执行增加超时控制、资源用量限制，"
    "以及更完善的错误信息透传机制，保障生产环境稳定性；"
)

set_para_text(paras[36],
    "·\xa0部署脚本完善：基于已有一键启动脚本设计，完成 dev/prod 双模式部署方案，"
    "撰写项目 README 中的部署与使用文档。"
)

new_doc.save(output_path)
print(f"Saved: {output_path}")
